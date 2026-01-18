import sys
# Python 3.13 တွင် cgi error မတက်စေရန် Patch ထည့်ခြင်း
try:
    import cgi
except ImportError:
    try:
        import legacy_cgi as cgi
        sys.modules['cgi'] = cgi
    except ImportError:
        pass

import streamlit as st
from googletrans import Translator
import PyPDF2
from docx import Document
from io import BytesIO
import time

# --- UI Configuration ---
st.set_page_config(page_title="Professional PDF Translator", page_icon="🌐")
st.markdown("""
    <style>
    .main { background-color: #fcfcfc; }
    .stButton>button { width: 100%; border-radius: 10px; height: 3em; font-weight: bold; }
    .status-info { background-color: #e1f5fe; padding: 15px; border-radius: 10px; border-left: 5px solid #01579b; }
    </style>
    """, unsafe_allow_html=True)

st.title("🌐 Smart PDF Translator (Auto-Resume)")
st.write("အင်္ဂလိပ်ဘာသာမှ မြန်မာဘာသာသို့ တစ်ကြောင်းချင်းစီ သေချာစွာ ပြန်ပေးပါသည်။")

# --- Session State (Resume စနစ်အတွက်) ---
if 'current_idx' not in st.session_state:
    st.session_state.current_idx = 0
if 'results' not in st.session_state:
    st.session_state.results = []
if 'working' not in st.session_state:
    st.session_state.working = False

translator = Translator()

uploaded_file = st.file_uploader("PDF ဖိုင်ကို ဤနေရာတွင် တင်ပါ", type="pdf")

if uploaded_file:
    reader = PyPDF2.PdfReader(uploaded_file)
    total_pages = len(reader.pages)
    
    st.markdown(f'<div class="status-info">စုစုပေါင်းစာမျက်နှာ: <b>{total_pages}</b> | လက်ရှိရောက်ရှိနေသောစာမျက်နှာ: <b>{st.session_state.current_idx}</b></div>', unsafe_allow_html=True)
    st.write("")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("▶️ ဘာသာပြန်ခြင်း စတင်/ဆက်လုပ်ရန်"):
            st.session_state.working = True
    with col2:
        if st.button("⏸️ ခေတ္တရပ်နားရန်"):
            st.session_state.working = False

    prog_bar = st.progress(st.session_state.current_idx / total_pages if total_pages > 0 else 0)
    log = st.empty()

    # --- ဘာသာပြန် လုပ်ငန်းစဉ် ---
    if st.session_state.working and st.session_state.current_idx < total_pages:
        for i in range(st.session_state.current_idx, total_pages):
            if not st.session_state.working:
                break
            
            try:
                page_text = reader.pages[i].extract_text()
                if page_text:
                    lines = page_text.split('\n')
                    translated_lines = []
                    
                    for line in lines:
                        if line.strip():
                            # တစ်ကြောင်းချင်းစီပြန်ခြင်း (Quality အတွက်)
                            res = translator.translate(line, src='en', dest='my')
                            translated_lines.append(res.text)
                            time.sleep(0.4) # Google Block မခံရအောင် နားချိန်ထည့်ခြင်း
                    
                    final_text = "\n".join(translated_lines)
                    st.session_state.results.append((f"Page {i+1}", final_text))
                    
                    st.session_state.current_idx = i + 1
                    prog_bar.progress(st.session_state.current_idx / total_pages)
                    log.success(f"✅ စာမျက်နှာ {i+1} ကို ဘာသာပြန်ဆိုပြီးပါပြီ။")
                
            except Exception as e:
                st.session_state.working = False
                st.error("⚠️ အင်တာနက်ပြတ်တောက်မှု သို့မဟုတ် ချိတ်ဆက်မှုပြဿနာ ဖြစ်ပေါ်နေပါသည်။ အင်တာနက်ပြန်ရလျှင် 'ဆက်လုပ်ရန်' ကို နှိပ်ပါ။")
                break

    # --- Word ဖိုင်အဖြစ် Download ရယူရန် ---
    if st.session_state.results:
        doc = Document()
        for page_num, text in st.session_state.results:
            doc.add_heading(page_num, level=2)
            doc.add_paragraph(text)
        
        output = BytesIO()
        doc.save(output)
        
        st.markdown("---")
        st.download_button(
            label="📥 ဘာသာပြန်ပြီးသမျှစာမျက်နှာများကို Word ဖိုင်အဖြစ် ရယူရန်",
            data=output.getvalue(),
            file_name="Translated_Document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
  )
      
